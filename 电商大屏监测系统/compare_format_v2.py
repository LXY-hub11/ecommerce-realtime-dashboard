# -*- coding: utf-8 -*-
"""Comprehensive format comparison between template and v2 generated report"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Emu

TEMPLATE = r"D:\项目\1\电商大屏监测系统\课程设计报告模板-大数据综合实训-0.docx"
OUTPUT = r"D:\项目\1\电商大屏监测系统\实时电商数据大屏监测系统_课程设计报告_v2.docx"

doc_t = Document(TEMPLATE)
doc_o = Document(OUTPUT)

def fmt_emu(val):
    if val is None: return 'None'
    emu = int(val)
    return f'{emu} EMU ({emu/12700:.1f}pt)'

def para_info(para):
    style = para.style.name if para.style else 'None'
    align = str(para.alignment) if para.alignment else 'None'
    pf = para.paragraph_format
    left_indent = pf.left_indent
    first_line_indent = pf.first_line_indent
    if para.runs:
        r = para.runs[0]
        r_font = r.font.name
        r_size = r.font.size
        r_bold = r.font.bold
    else:
        r_font = r_size = r_bold = None
    return {
        'style': style, 'align': align,
        'left_indent': left_indent, 'first_line_indent': first_line_indent,
        'font': r_font, 'size': r_size, 'bold': r_bold,
    }

print("=" * 130)
print("COMPREHENSIVE FORMAT COMPARISON: Template (T) vs Generated v2 (O)")
print("=" * 130)

# Count format issues
format_issues = []
all_match = True

# Check P020 onwards
for i in range(20, max(len(doc_t.paragraphs), len(doc_o.paragraphs))):
    if i >= len(doc_t.paragraphs):
        print(f"P{i:03d} [EXTRA in output]")
        all_match = False
        continue
    if i >= len(doc_o.paragraphs):
        print(f"P{i:03d} [MISSING in output]")
        all_match = False
        continue

    pt = doc_t.paragraphs[i]
    po = doc_o.paragraphs[i]
    it = para_info(pt)
    io = para_info(po)

    issues = []
    if it['style'] != io['style']:
        issues.append(f"STYLE: {it['style']} → {io['style']}")
    if it['font'] != io['font']:
        issues.append(f"FONT: {it['font']} → {io['font']}")
    if it['size'] != io['size']:
        issues.append(f"SIZE: {fmt_emu(it['size'])} → {fmt_emu(io['size'])}")
    if it['bold'] != io['bold']:
        issues.append(f"BOLD: {it['bold']} → {io['bold']}")
    if it['first_line_indent'] != io['first_line_indent']:
        issues.append(f"1stINDENT: {fmt_emu(it['first_line_indent'])} → {fmt_emu(io['first_line_indent'])}")
    if it['left_indent'] != io['left_indent']:
        issues.append(f"LEFT: {fmt_emu(it['left_indent'])} → {fmt_emu(io['left_indent'])}")
    if it['align'] != io['align']:
        issues.append(f"ALIGN: {it['align']} → {io['align']}")

    if issues:
        format_issues.append((i, issues))
        all_match = False
        status = "⚠"
    else:
        status = "✓"

    # Only print non-matching lines
    if issues:
        text_preview = po.text[:60] if po.text else '(empty)'
        print(f"P{i:03d} {status} [{it['style'][:26]:<26}] font={str(it['font']):<18} size={fmt_emu(it['size']):<20} "
              f"bold={str(it['bold']):<5} 1st={fmt_emu(it['first_line_indent']):<20} "
              f"left={fmt_emu(it['left_indent']):<18} align={it['align']} →\"{text_preview}\"")
        for iss in issues:
            print(f"      → {iss}")

# Summary
print("\n" + "=" * 130)
if all_match:
    print("✅ ALL FORMATTING MATCHES TEMPLATE EXACTLY!")
else:
    print(f"⚠ FORMAT ISSUES FOUND: {len(format_issues)}")
    for idx, issues in format_issues:
        print(f"\nP{idx:03d}:")
        print(f"  Template: [{para_info(doc_t.paragraphs[idx])['style']}] \"{doc_t.paragraphs[idx].text[:80]}...\"")
        po = doc_o.paragraphs[idx]
        print(f"  Output:   [{para_info(po)['style']}] \"{po.text[:80]}...\"")
        for iss in issues:
            print(f"  → {iss}")

# Also verify pages 1-2 are untouched
print("\n--- Pages 1-2 Verification ---")
pages_ok = True
for i in range(20):
    tt = doc_t.paragraphs[i].text
    to = doc_o.paragraphs[i].text
    if tt != to:
        print(f"  ⚠ P{i:03d} CHANGED: \"{tt[:50]}\" → \"{to[:50]}\"")
        pages_ok = False
if pages_ok:
    print("  ✅ Pages 1-2 (P000-P019) completely unchanged")

print("\nDone.")
