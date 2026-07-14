# -*- coding: utf-8 -*-
"""Intelligent verification with correct shift"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

TEMPLATE = r"D:\项目\1\电商大屏监测系统\课程设计报告模板-大数据综合实训-0.docx"
OUTPUT = r"D:\项目\1\电商大屏监测系统\实时电商数据大屏监测系统_课程设计报告_v2.docx"

doc_t = Document(TEMPLATE)
doc_o = Document(OUTPUT)

SHIFT = 9  # 9 refs [7]-[15] inserted before appendix

def fmt_emu(val):
    if val is None: return 'None'
    return f'{int(val)} EMU ({int(val)/12700:.1f}pt)'

def para_fmt(para):
    style = para.style.name if para.style else 'None'
    pf = para.paragraph_format
    align = str(para.alignment) if para.alignment else 'None'
    if para.runs:
        r = para.runs[0]
        return {
            'style': style, 'align': align,
            'size': r.font.size, 'font': r.font.name, 'bold': r.font.bold,
            'first_line_indent': pf.first_line_indent, 'left_indent': pf.left_indent,
        }
    return {
        'style': style, 'align': align,
        'size': None, 'font': None, 'bold': None,
        'first_line_indent': pf.first_line_indent, 'left_indent': pf.left_indent,
    }

def check_match(t_idx, o_idx, label):
    """Compare formatting of template[t_idx] vs output[o_idx]"""
    ft = para_fmt(doc_t.paragraphs[t_idx])
    fo = para_fmt(doc_o.paragraphs[o_idx])
    mismatches = []
    for key in ['style', 'align', 'size', 'font', 'bold', 'first_line_indent', 'left_indent']:
        if ft[key] != fo[key]:
            mismatches.append(f"  {key}: {ft[key]} -> {fo[key]}")
    if mismatches:
        tt = doc_t.paragraphs[t_idx].text[:60]
        ot = doc_o.paragraphs[o_idx].text[:60]
        print(f"  FAIL T{t_idx:03d}->O{o_idx:03d} [{ft['style']}]")
        print(f"    Template: \"{tt}\"")
        print(f"    Output:   \"{ot}\"")
        for m in mismatches:
            print(m)
        return False
    else:
        tt = doc_t.paragraphs[t_idx].text[:50]
        print(f"  OK   T{t_idx:03d}->O{o_idx:03d} [{ft['style']}] \"{tt}...\"")
        return True

print("=" * 100)
print("FINAL FORMAT VERIFICATION")
print(f"SHIFT = {SHIFT} paragraphs inserted before appendix")
print("=" * 100)

all_ok = True

# 1. P020-P213: direct match (no shift)
print("\n--- P020-P213 (no shift) ---")
for i in range(20, 214):
    if not check_match(i, i, f'P{i:03d}'):
        all_ok = False

# 2. Inserted refs P214-P222: should match ref format from P208
print("\n--- Inserted refs P214-P222 ---")
ref_fmt = para_fmt(doc_t.paragraphs[208])
for i in range(214, 223):
    fo = para_fmt(doc_o.paragraphs[i])
    ok = True
    for key in ['style', 'size', 'bold']:
        if ref_fmt[key] != fo[key]:
            print(f"  FAIL P{i:03d} {key}: expected {ref_fmt[key]}, got {fo[key]}")
            ok = False
            all_ok = False
    if ok:
        print(f"  OK   P{i:03d} [{fo['style']}] \"{doc_o.paragraphs[i].text[:70]}\"")

# 3. Shifted appendix: T214-T218 -> O(214+9)-O(218+9) = O223-O227
print(f"\n--- Shifted appendix T214-T218 -> O{214+SHIFT}-O{218+SHIFT} ---")
for t_idx in range(214, 219):
    o_idx = t_idx + SHIFT
    if not check_match(t_idx, o_idx, f'T{t_idx:03d}->O{o_idx:03d}'):
        all_ok = False

# 4. Pages 1-2
print("\n--- Pages 1-2 ---")
p12_ok = True
for i in range(20):
    if doc_t.paragraphs[i].text != doc_o.paragraphs[i].text:
        print(f"  FAIL P{i:03d} changed!")
        p12_ok = False
if p12_ok:
    print("  OK - Completely unchanged")

# 5. Additional output paragraphs (beyond template range)
print("\n--- Extra paragraphs ---")
if len(doc_o.paragraphs) > 219 + SHIFT:
    for i in range(219 + SHIFT, len(doc_o.paragraphs)):
        print(f"  EXTRA P{i:03d}: [{para_fmt(doc_o.paragraphs[i])['style']}] \"{doc_o.paragraphs[i].text[:60]}\"")

# Summary
print("\n" + "=" * 100)
if all_ok:
    print("ALL FORMATTING MATCHES PERFECTLY!")
else:
    print("Some format issues found - see details above.")
print("=" * 100)
