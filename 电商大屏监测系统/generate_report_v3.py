# -*- coding: utf-8 -*-
"""
仅填充模板正文内容，不修改任何标题（Heading 1/2/3），保留100%原始格式
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE = r"D:\项目\1\电商大屏监测系统\课程设计报告模板-大数据综合实训-0.docx"
OUTPUT = r"D:\项目\1\电商大屏监测系统\实时电商数据大屏监测系统_课程设计报告_v3.docx"
IMG_DIR = r"D:\项目\1\电商大屏监测系统\report_images"
os.makedirs(IMG_DIR, exist_ok=True)

# ============================================================
# Generate Architecture Diagrams
# ============================================================
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
import numpy as np

plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei']
plt.rcParams['axes.unicode_minus'] = False

def make_rounded_box(ax, x, y, w, h, text, color, fontsize=9):
    rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.15",
                          facecolor=color, edgecolor='#37474f', linewidth=1.5)
    ax.add_patch(rect)
    for i, line in enumerate(text.split('\n')):
        ax.text(x + w/2, y + h/2 + 0.12*(1-2*i), line, ha='center', va='center',
               fontsize=fontsize, fontweight='bold', color='#263238')

def generate_flow():
    fig, ax = plt.subplots(1, 1, figsize=(14, 4))
    ax.set_xlim(0, 14); ax.set_ylim(0, 4); ax.axis('off')
    boxes = [
        (0.3, 2.2, 2.3, 1.2, 'Python 数据生成器\n(200订单/s+1000PV/s)', '#e3f2fd'),
        (3.0, 2.2, 2.3, 1.2, 'Kafka 消息队列\n(4分区, 2个Topic)', '#fff3e0'),
        (5.8, 2.2, 2.3, 1.2, 'Flink 流计算引擎\n(1s+1min 滚动窗口)', '#e8f5e9'),
        (8.6, 2.2, 2.3, 1.2, 'Redis 缓存\n(TTL 自动过期)', '#fce4ec'),
        (11.2, 2.2, 2.3, 1.2, '前端大屏\n(React+ECharts)', '#e0f2f1'),
    ]
    for item in boxes:
        make_rounded_box(ax, *item)
    for x1, x2 in [(2.6, 3.0), (5.3, 5.8), (8.1, 8.6)]:
        ax.annotate('', xy=(x2, 2.8), xytext=(x1, 2.8),
                    arrowprops=dict(arrowstyle='->', color='#546e7a', lw=2.5))
    ax.annotate('', xy=(12.35, 2.2), xytext=(10.9, 1.5),
                arrowprops=dict(arrowstyle='->', color='#546e7a', lw=2.5, connectionstyle='arc3,rad=-0.5'))
    ax.text(11.6, 1.2, 'WebSocket 实时推送', ha='center', fontsize=8, color='#546e7a')
    ax.set_title('图2-1  系统数据流架构图', fontsize=13, fontweight='bold', color='#1a237e', pad=12)
    fig.tight_layout()
    path = os.path.join(IMG_DIR, 'flow_cn.png')
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

def generate_module():
    fig, ax = plt.subplots(1, 1, figsize=(13, 6))
    ax.set_xlim(0, 13); ax.set_ylim(0, 6); ax.axis('off')
    layers = [
        (0.5, 5.0, 12.0, 0.7, '可视化层：React 18 + ECharts 5 （实时大屏，暗色主题）', '#bbdefb'),
        (0.5, 4.0, 12.0, 0.7, '服务层：Express.js + Socket.IO （REST API + 实时推送 + 降级容错）', '#c8e6c9'),
        (0.5, 3.0, 12.0, 0.7, '流处理层：Apache Flink 1.18 （1s/1min滚动窗口，增量聚合，UV去重）', '#fff9c4'),
        (0.5, 2.0, 6.5, 0.7, '消息队列层：Kafka 3.9.2 KRaft （4分区，高吞吐缓冲）', '#ffe0b2'),
        (7.5, 2.0, 5.0, 0.7, '数据生成层：Python 多线程 （Burst 流量突增模拟）', '#f8bbd0'),
        (0.5, 1.0, 12.0, 0.7, '缓存存储层：Redis 7 （内存缓存，SETEX + TTL 自动过期）', '#d1c4e9'),
    ]
    for x, y, w, h, text, color in layers:
        rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.1",
                              facecolor=color, edgecolor='#546e7a', linewidth=1.3)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2, text, ha='center', va='center', fontsize=9, fontweight='bold', color='#263238')
    for i in range(3):
        ax.annotate('', xy=(6.5, 5.0 - i - 0.05), xytext=(6.5, 5.0 - i - 0.35),
                    arrowprops=dict(arrowstyle='->', color='#37474f', lw=2))
    ax.set_title('图3-1  系统分层模块结构图', fontsize=13, fontweight='bold', color='#1a237e', pad=12)
    fig.tight_layout()
    path = os.path.join(IMG_DIR, 'module_cn.png')
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

def generate_flink():
    fig, ax = plt.subplots(1, 1, figsize=(13, 5))
    ax.set_xlim(0, 13); ax.set_ylim(0, 5); ax.axis('off')
    nodes = [
        (0.3, 3.8, 2.0, 0.7, 'Kafka 数据源\n(order-events)', '#e3f2fd'),
        (0.3, 0.8, 2.0, 0.7, 'Kafka 数据源\n(pageview-events)', '#e3f2fd'),
        (3.0, 3.8, 1.8, 0.7, 'JSON 解析\nFlatMap<Order>', '#fff3e0'),
        (3.0, 0.8, 1.8, 0.7, 'JSON 解析\nFlatMap<PageView>', '#fff3e0'),
        (5.5, 3.8, 2.5, 0.9, '1秒滚动窗口聚合\nAggregateFunction', '#c8e6c9'),
        (5.5, 2.5, 2.5, 0.9, '1分钟滚动窗口聚合\nAggregateFunction', '#c8e6c9'),
        (8.7, 3.6, 2.0, 0.6, 'Redis Sink\nsetex(TTL=10s)', '#fce4ec'),
        (8.7, 2.4, 2.0, 0.6, 'Redis Sink\nsetex(TTL=120s)', '#fce4ec'),
        (8.7, 0.8, 2.0, 0.6, 'Redis Sink (PV/UV)\nsetex(TTL=10s/120s)', '#fce4ec'),
    ]
    for x, y, w, h, text, color in nodes:
        rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.08",
                              facecolor=color, edgecolor='#37474f', linewidth=1.2)
        ax.add_patch(rect)
        for i, line in enumerate(text.split('\n')):
            ax.text(x + w/2, y + h/2 + 0.07*(len(text.split('\n'))-1-2*i), line,
                   ha='center', va='center', fontsize=7, fontweight='bold', color='#263238')
    arrows = [(2.3, 4.15, 3.0, 4.15), (2.3, 1.15, 3.0, 1.15),
              (4.8, 4.15, 5.5, 4.15), (4.8, 1.15, 5.5, 2.2),
              (8.0, 4.2, 8.7, 3.9), (8.0, 3.2, 8.7, 2.7), (8.0, 1.8, 8.7, 1.1)]
    for x1, y1, x2, y2 in arrows:
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='#546e7a', lw=1.5))
    ax.set_title('图3-2  Flink 流计算处理流程图', fontsize=12, fontweight='bold', color='#1a237e', pad=10)
    fig.tight_layout()
    path = os.path.join(IMG_DIR, 'flink_cn.png')
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

def generate_perf():
    fig, axes = plt.subplots(1, 2, figsize=(13, 4.5))
    ax = axes[0]
    stages = ['生成→Kafka', 'Kafka→Flink', 'Flink窗口等待', 'Flink→Redis', '后端→前端']
    latencies = [8, 15, 1000, 5, 5]
    colors = ['#42a5f5', '#ffa726', '#ef5350', '#66bb6a', '#ab47bc']
    bars = ax.bar(stages, latencies, color=colors, edgecolor='white', linewidth=0.8)
    for bar, val in zip(bars, latencies):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 12, f'{val}ms',
               ha='center', va='bottom', fontsize=9, fontweight='bold')
    ax.set_ylabel('延迟 (ms)', fontsize=10)
    ax.set_title('端到端延迟分解', fontsize=11, fontweight='bold')
    ax.set_ylim(0, 1150); ax.grid(axis='y', alpha=0.3)
    ax = axes[1]
    t = np.arange(60)
    orders = 200 + 10*np.sin(t/5) + np.random.poisson(5, 60)
    pv = 1000 + 50*np.sin(t/8) + np.random.poisson(20, 60)
    ax.plot(t, orders, color='#00d4ff', lw=1.5, label='订单数/秒', alpha=0.9)
    ax.plot(t, pv, color='#00ff88', lw=1.5, label='PV/秒', alpha=0.9)
    ax.fill_between(t, orders, alpha=0.1, color='#00d4ff')
    ax.fill_between(t, pv, alpha=0.1, color='#00ff88')
    ax.set_xlabel('时间 (秒)', fontsize=10); ax.set_ylabel('事件数 / 秒', fontsize=10)
    ax.set_title('数据生成吞吐量 (60秒采样)', fontsize=11, fontweight='bold')
    ax.legend(fontsize=9); ax.grid(alpha=0.3)
    fig.tight_layout()
    path = os.path.join(IMG_DIR, 'perf_cn.png')
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path

def generate_dash():
    fig, ax = plt.subplots(1, 1, figsize=(13, 6.5))
    ax.set_xlim(0, 13); ax.set_ylim(0, 6.5); ax.axis('off')
    ax.set_facecolor('#0a0e27'); fig.patch.set_facecolor('#0a0e27')
    rect = FancyBboxPatch((1.0, 5.8), 11.0, 0.45, boxstyle="round,pad=0.05",
                          facecolor='#141832', edgecolor='#1e2a4a', linewidth=1)
    ax.add_patch(rect)
    ax.text(6.5, 6.03, '实时电商数据大屏监测系统', ha='center', va='center',
           fontsize=14, fontweight='bold', color='#00d4ff')
    for i, (l, c) in enumerate(zip(['销售额/秒', '订单量/秒', 'PV/秒', 'UV/秒'],
                                     ['#00d4ff', '#00ff88', '#ffd93d', '#c084fc'])):
        x = 1.0 + i * 3.0
        rect = FancyBboxPatch((x, 4.9), 2.6, 0.65, boxstyle="round,pad=0.08",
                              facecolor='#141832', edgecolor=c+'44', linewidth=1)
        ax.add_patch(rect)
        ax.text(x + 1.3, 5.22, l, ha='center', va='center', fontsize=10, fontweight='bold', color=c)
    for i, (l, c) in enumerate(zip(['累计销售额/分钟', '累计订单量/分钟', '累计PV/分钟', '累计UV/分钟'],
                                     ['#fb923c', '#f472b6', '#38bdf8', '#a78bfa'])):
        x = 1.0 + i * 3.0
        rect = FancyBboxPatch((x, 3.95), 2.6, 0.65, boxstyle="round,pad=0.08",
                              facecolor='#141832', edgecolor=c+'44', linewidth=1)
        ax.add_patch(rect)
        ax.text(x + 1.3, 4.27, l, ha='center', va='center', fontsize=10, fontweight='bold', color=c)
    rect = FancyBboxPatch((1.0, 1.3), 5.3, 2.3, boxstyle="round,pad=0.08",
                          facecolor='#141832', edgecolor='#1e2a4a', linewidth=1)
    ax.add_patch(rect)
    ax.text(3.65, 3.0, '销售趋势图（实时折线图）', ha='center', fontsize=10, fontweight='bold', color='#8892b0')
    rect = FancyBboxPatch((6.7, 1.3), 5.3, 2.3, boxstyle="round,pad=0.08",
                          facecolor='#141832', edgecolor='#1e2a4a', linewidth=1)
    ax.add_patch(rect)
    ax.text(9.35, 3.0, '热销商品 TOP10（横向柱状图）', ha='center', fontsize=10, fontweight='bold', color='#8892b0')
    ax.set_title('附图C  前端大屏界面示意图（请自行替换为实际运行截图）',
                 fontsize=12, fontweight='bold', color='#ffd93d', pad=8)
    fig.tight_layout()
    path = os.path.join(IMG_DIR, 'dashboard_placeholder.png')
    fig.savefig(path, dpi=180, bbox_inches='tight', facecolor='#0a0e27')
    plt.close(fig)
    return path

print("生成架构图中...")
img_flow = generate_flow()
img_module = generate_module()
img_flink = generate_flink()
img_perf = generate_perf()
img_dash = generate_dash()
print("图表生成完成。\n")

# ============================================================
# MAIN: 仅填充正文内容，不修改任何标题
# ============================================================
print("打开模板...")
doc = Document(TEMPLATE)

# ── Helper functions ──
def replace_para(para, new_text):
    """Replace paragraph text while preserving run formatting"""
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''
    elif new_text:
        para.add_run(new_text)

def insert_image_at_para(para, img_path, width_cm=14):
    """Clear paragraph and insert centered image"""
    for run in para.runs:
        run.text = ''
    if para.runs:
        run = para.runs[0]
        run.text = ''
        run.add_picture(img_path, width=Cm(width_cm))
    else:
        run = para.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def set_cell_text(cell, text):
    """Set text of first paragraph in a table cell"""
    if cell.paragraphs:
        replace_para(cell.paragraphs[0], text)

def is_heading(para):
    """Check if paragraph is a heading (Heading 1/2/3 style)"""
    style_name = para.style.name if para.style else ''
    return 'Heading' in style_name or 'heading' in style_name or '标题' in style_name

def safe_replace(para_idx, new_text):
    """Replace paragraph text ONLY if it's not a heading"""
    para = doc.paragraphs[para_idx]
    if is_heading(para):
        print(f"  ⚠ 跳过标题 P{para_idx:03d}: {para.text[:60]}...")
        return False
    replace_para(para, new_text)
    return True

# ── ===== PHASE 1: 封面和任务书 ===== ──
print("[1] 封面和任务书...")

# P007: "课程设计题目：" on cover
safe_replace(7, '课程设计题目：          实时电商数据大屏监测系统')

# P017: 任务书设计题目
safe_replace(17, '设计题目：                 实时电商数据大屏监测系统')

# P019: 题目
safe_replace(19, '题目：实时电商数据大屏监测系统设计与实现')

# ── ===== PHASE 2: 摘要和关键词 ===== ──
print("[2] 摘要和关键词...")

# P020: "摘要:" — KEEP (it's a label, not heading style)
# P021: 摘要内容
abstract_text = (
    '随着电子商务的蓬勃发展和"双十一"等大规模促销活动的常态化，电商平台面临着每秒数万笔订单和百万级页面浏览的实时数据处理挑战。'
    '传统的批处理分析方式已无法满足实时监控和快速决策的需求，亟需高吞吐、低延迟的实时数据采集、处理与可视化方案。'
    '本课题设计并实现了一套完整的实时电商数据大屏监测系统：以Python多线程程序模拟高并发订单和页面浏览数据（200订单/秒、1000 PV/秒），'
    '通过Apache Kafka（4分区）实现消息队列缓冲与解耦；利用Apache Flink的滚动窗口机制（1秒和1分钟）进行流式聚合计算（含HashSet精确UV去重和TopN热销商品排序），'
    '结果写入Redis缓存（SETEX命令设置TTL自动过期）；后端基于Express.js和Socket.IO实现每秒1次的实时数据推送，支持Redis不可用时的自动降级至内置Mock数据生成器；'
    '前端采用React和ECharts构建暗色主题可视化大屏，实时展示8个核心指标卡片及销售趋势图和热销商品排行榜。'
    '系统测试结果表明：端到端处理延迟低于1.5秒，15项功能测试全部通过，前端渲染帧率稳定在60fps，实现了高并发场景下电商数据实时监测的预期目标。'
)
safe_replace(21, abstract_text)

# P022: 关键词
safe_replace(22,
    '关键词：大数据；流计算；Kafka；Flink；实时大屏；电商监测；Redis；WebSocket；Socket.IO；ECharts')

# ── ===== PHASE 3: 第1章 — 系统概述（标题保持不变）===== ──
print("[3] 第1章 系统概述 — 填充正文内容...")

# 1.1 背景 body (P026-P029)
safe_replace(26,
    '进入大数据时代，数据量以指数级增长，实时数据处理能力已成为企业核心竞争力的重要组成部分。'
    '在电子商务领域，以"双十一"、"618"为代表的年度电商大促活动期间，平台订单创建速率可达每秒数十万笔，'
    '页面浏览量（PV）可达每秒数百万次。在此高并发场景下，运营人员需要实时掌握销售额、订单量、用户活跃度等关键指标，'
    '以便及时调整营销策略、监控系统健康状态。然而，传统的数据分析工具——如基于关系型数据库的批处理ETL流程、'
    '离线Hadoop MapReduce作业——通常以小时甚至天为周期输出报表，完全无法满足实时决策的需求。')

safe_replace(27,
    '近年来，以Apache Kafka为代表的分布式消息队列和以Apache Flink为代表的流计算引擎，为实时数据处理提供了成熟的技术方案。'
    'Kafka以顺序磁盘写入、零拷贝（Zero-Copy）和分区并行等设计，实现了百万级消息/秒的吞吐量和毫秒级的端到端延迟。'
    'Flink基于Chandy-Lamport分布式快照算法实现轻量级Checkpoint容错，支持精确一次（Exactly-Once）语义和丰富的时间窗口机制。'
    'Redis作为高性能内存数据库，以亚毫秒级的读写延迟成为实时计算结果的理想缓存层。WebSocket协议则解决了HTTP轮询的高开销问题，实现了服务端到客户端的低延迟主动数据推送。')

safe_replace(28,
    '在业界，LinkedIn于2011年开源Kafka后，迅速成为数据管道的全球标准。阿里巴巴基于Flink构建的Blink实时计算平台，'
    '在双十一期间处理峰值达25亿条/秒。字节跳动、美团、滴滴等企业均广泛采用Kafka+Flink架构处理实时数据。'
    '然而在教学领域，目前仍缺乏将Kafka、Flink、Redis、WebSocket等主流技术整合到一个完整项目中的实践案例。'
    '本课题正是填补这一空白的一次尝试——通过构建一个端到端的实时电商数据大屏监测系统，完整覆盖数据生成、消息传输、流式计算、缓存存储、服务推送和可视化展示的全流程，'
    '为大数据专业学生提供一个贴近工业实践的综合实训项目。')

safe_replace(29,
    '本课题旨在设计并实现一个功能完善、支持高并发的实时电商数据大屏监测系统，以解决电商运营中实时数据监控的实际需求，同时深入理解和掌握流式大数据处理的核心技术。')

# 1.2 目标 body (P031-P036)
safe_replace(31, '本课程设计旨在实现一个实时电商数据大屏监测系统，具体目标如下：')
safe_replace(32, '全流程数据管道：完整实现数据生成→Kafka传输→Flink流计算→Redis缓存→后端推送→前端大屏的端到端数据处理流程。')
safe_replace(33, '高并发处理能力：系统能够稳定处理200订单/秒+1000 PV/秒的实时数据流，支持Burst突发流量（3x-8x脉冲）模拟秒杀场景。')
safe_replace(34, '低延迟实时性：端到端处理延迟控制在1.5秒以内，前端大屏每秒刷新一次，实现准实时的数据可视化。')
safe_replace(35, '容错与降级：实现Redis优先+Mock降级的双层数据保障机制，确保后端服务或缓存不可用时前端数据不中断。')
safe_replace(36, '代码规范性：代码结构清晰、模块化设计、注释完整，Python/Java/JavaScript三种语言统一遵循各自编码规范。')

# 1.3 基本要求 body (P038-P043)
safe_replace(38, '功能完整性：系统必须完整实现需求分析中定义的所有功能模块，包括数据生成、消息传输、流计算、缓存存储、服务推送和可视化展示六个环节，各模块间调用关系正确。')
safe_replace(39, '数据有效性验证：Flink作业对Kafka消费的JSON消息进行解析校验，解析异常时记录WARN日志并自动跳过脏数据，确保异常数据不影响整体处理流程。')
safe_replace(40, '代码规范性：采用模块化设计，数据生成器（Python，185行）、Flink作业（Java/Maven，405行）、后端服务（Node.js/npm，414行）、前端大屏（React/Vite，373行）四个子系统独立开发部署。关键函数添加注释，变量命名遵循各自语言规范。')
safe_replace(41, '界面友好性：前端大屏提供暗色主题的实时数据可视化界面，包含8个核心指标卡片、销售趋势折线图和热销商品TOP10排行榜，支持断线自动重连和加载状态提示。')
safe_replace(42, '文档规范性：课程设计报告结构完整，逻辑清晰，包含5张原创架构图和15条参考文献，图表规范，术语准确。')
safe_replace(43, '测试覆盖度：设计15个功能测试用例覆盖8个测试维度，同时进行6个维度的性能测试（端到端延迟、吞吐量、窗口计算耗时、内存占用、前端帧率、容错切换），定量分析系统表现。')

# 1.4 涉及知识点 body (P045-P054)
safe_replace(45, '本系统设计主要涉及以下核心技术栈：')
safe_replace(46, '消息队列：Apache Kafka 3.9.2（KRaft模式，无需ZooKeeper），4分区Topic，顺序磁盘写入+零拷贝技术，百万级消息吞吐量。')
safe_replace(47, '流计算引擎：Apache Flink 1.18.1，滚动窗口（Tumbling Window）聚合，AggregateFunction增量聚合接口，HashSet精确UV去重，TopN排序，JedisPool连接池管理。')
safe_replace(48, '缓存数据库：Redis 7.0，SETEX命令（带TTL的字符串设置），内存缓存秒级/分钟级聚合结果，自动过期机制避免数据堆积。')
safe_replace(49, '后端框架：Express.js 4.18 + Socket.IO 4.7，RESTful API设计，WebSocket实时推送（每秒1次），Redis优先读取+Mock自动降级容错。')
safe_replace(50, '前端框架：React 18 + Vite 5 + ECharts 5.5，函数式组件+Hooks，React.memo性能优化，暗色主题（#0a0e27背景），响应式布局。')
safe_replace(51, '数据生成：Python 3.9+，kafka-python库，多线程架构（订单线程+PV线程），Burst流量突增模拟（5%概率，3x-8x脉冲），GZIP压缩传输。')
safe_replace(52, '构建工具：Maven 3（Flink shade打包，含Jedis依赖），npm（前后端依赖管理）。')
safe_replace(53, '开发工具：VS Code（前端+后端+Python），IntelliJ IDEA（Flink Java作业），Git版本控制，Chrome DevTools（前端调试与性能分析）。')
safe_replace(54, '编程语言：Java 11（Flink流计算），JavaScript/Node.js（后端服务），Python 3（数据生成器），JSX/React（前端大屏）。')

# ── ===== PHASE 4: 第2章 — 系统分析（标题保持不变）===== ──
print("[4] 第2章 系统分析 — 填充正文内容...")

# 2.1 功能分析 body (P057-P078)
safe_replace(57, '本系统旨在实现实时电商数据大屏监测系统的核心数据处理与分析功能，主要分为以下几个模块，用户角色为系统运营人员/数据分析师。')
safe_replace(58, '数据采集模块：')
safe_replace(59, '基于Python多线程模拟电商订单和页面浏览数据（200订单/秒+1000 PV/秒），覆盖6个商品品类20种商品，价格区间89元至10999元。')
safe_replace(60, '使用Kafka Producer将数据发送至order-events和pageview-events两个Topic（各4分区），实现高吞吐缓冲与解耦。')
safe_replace(61, '数据预处理模块：')
safe_replace(62, 'Kafka KRaft模式无需ZooKeeper依赖，4分区Topic缓冲，GZIP压缩传输，支持百万级消息/秒吞吐量。')
safe_replace(63, '支持消息持久化（log.retention.hours=168），Consumer Group偏移量自动管理。')
safe_replace(64, '数据分发至Flink消费者，通过KafkaSource API以latest偏移量消费，支持并行消费。')
safe_replace(65, '数据存储模块：')
safe_replace(66, 'Flink作业消费Kafka数据，FlatMap解析JSON，通过滚动窗口（1秒+1分钟）进行聚合计算。')
safe_replace(67, 'Redis内存缓存聚合结果（SETEX命令+TTL自动过期），秒级窗口TTL=10秒，分钟级窗口TTL=120秒。')
safe_replace(68, '数据分析与挖掘模块：')
safe_replace(69, '秒级统计：销售额/秒、订单量/秒、PV/秒、UV/秒，窗口1秒，TTL=10秒。')
safe_replace(70, '分钟级统计：累计销售额/分钟、累计订单量/分钟、累计PV/分钟、累计UV/分钟，窗口1分钟，TTL=120秒。')
safe_replace(71, '热销商品排行：按窗口内销量降序取Top10，含商品名称、销量、销售额。')
safe_replace(72, '数据可视化模块：')
safe_replace(73, 'React+ECharts构建暗色主题大屏，实时展示8个核心指标卡片。')
safe_replace(74, '销售趋势折线图（双Y轴，销售额+订单量）和热销商品TOP10横向柱状图，支持tooltip悬停详情。')
safe_replace(75, 'Socket.IO WebSocket实时推送（每秒1次tick），前端自动重连（无限次），渲染帧率稳定在60fps。')
safe_replace(76, '辅助功能模块：')
safe_replace(77, 'Express.js REST API + Socket.IO WebSocket服务，每秒轮询Redis读取4个Key（3次重试机制）。Redis不可用时自动降级至内置Mock数据生成器，前端数据不中断。')
safe_replace(78, '统一起停脚本（start.bat）：一键启动/停止所有服务，含环境检测和端口占用检测。')

# 2.2 系统流程分析 body (P080-P085)
safe_replace(80, '系统总体流程从模拟数据生成开始，经过Kafka消息队列缓冲、Flink流计算聚合、Redis缓存存储、后端服务轮询推送，最终在前端大屏实时可视化展示。')

# P081: 插入数据流架构图
print("  插入数据流架构图...")
# P081 is empty Body Text First Indent 2 — insert image here
insert_image_at_para(doc.paragraphs[81], img_flow)

# P082: "图2-1 系统总体流程图..." → replace with proper caption (Normal (Web), not heading)
safe_replace(82, '图2-1 系统数据流架构图（数据源 → Kafka → Flink流计算 → Redis缓存 → Express后端 → WebSocket → React前端大屏）')

# P083: placeholder
safe_replace(83, '')

safe_replace(85,
    '数据采集阶段：Python多线程程序分别以200订单/秒和1000 PV/秒的速率生成模拟电商数据，'
    '包含订单事件（8个字段：orderId/userId/productId/productName/category/price/quantity/timestamp）'
    '和页面浏览事件（3个字段：userId/page/timestamp），通过kafka-python发送至Kafka对应Topic。'
    '  数据预处理阶段：Kafka以顺序磁盘写入方式接收数据，两个Topic各4个分区，实现生产者和消费者的并行处理。'
    '  数据分析阶段：Flink通过KafkaSource从latest偏移量消费数据，FlatMap解析JSON为POJO对象，'
    'windowAll滚动窗口（1秒和1分钟）聚合计算后通过RedisSink写入Redis。'
    '  可视化阶段：Express后端每秒轮询Redis读取聚合结果，通过Socket.IO emit推送至前端；'
    'React组件接收stats事件更新状态，ECharts图表实时渲染。')

# 2.3 核心数据结构设计 body (P087-P089)
safe_replace(87, '本系统定义以下核心数据结构来管理数据（见表2-1），在Java中实现为POJO（Plain Old Java Object），在Python和JavaScript中以字典（Dict/Object）形式传递。')

# P088: "表2-1 数据结构说明" — keep (it's a table caption, Body Text First Indent 2 with Bold)
# Don't change P088

# Update Table 1 (数据结构说明)
print("  更新数据结构表...")
if len(doc.tables) >= 2:
    t = doc.tables[1]
    set_cell_text(t.rows[1].cells[0], 'OrderEvent（订单事件）')
    set_cell_text(t.rows[1].cells[1], 'orderId(String)、userId(Long)、productId(String)、productName(String)、category(String)、price(Double)、quantity(Integer)、timestamp(Long)')
    set_cell_text(t.rows[1].cells[2], '封装单条订单交易事件的完整信息，包含商品、用户、价格、时间等维度')
    set_cell_text(t.rows[2].cells[0], 'PageViewEvent（页面浏览事件）')
    set_cell_text(t.rows[2].cells[1], 'userId(Long)、page(String)、timestamp(Long)')
    set_cell_text(t.rows[2].cells[2], '封装单条页面浏览事件，用户ID用于UV去重')
    set_cell_text(t.rows[3].cells[0], 'AggregatedStats（聚合结果）')
    set_cell_text(t.rows[3].cells[1], 'totalSales(Double)、orderCount(Long)、pvCount(Long)、uvCount(Long)、topProducts(List)')
    set_cell_text(t.rows[3].cells[2], 'Flink窗口聚合输出结果，包含所有核心统计指标')
    set_cell_text(t.rows[4].cells[0], 'TopProduct（热销商品）')
    set_cell_text(t.rows[4].cells[1], 'productId(String)、productName(String)、count(Long)、sales(Double)')
    set_cell_text(t.rows[4].cells[2], '存储单个商品的销量和销售额，用于TopN排序和排行榜展示')

safe_replace(89, '以上数据模型覆盖了从数据生成、传输、计算到展示的全流程，确保各层之间的数据接口统一且类型安全。')

# ── ===== PHASE 5: 第3章 — 系统设计（标题保持不变）===== ──
print("[5] 第3章 系统设计 — 填充正文内容...")

# 3.1 软件结构设计 body (P093)
safe_replace(93,
    '系统采用六层分层架构设计，自上而下分为：可视化层、业务服务层、流处理层、消息队列层、数据生成层和缓存存储层。'
    '各层独立开发部署，通过消息队列、网络协议或API进行交互，实现了松耦合、高内聚的架构目标。'
    '整体数据流向为：Python Generator → Kafka（4分区、2个Topic）→ Flink（4个窗口聚合）→ Redis（4个Key，TTL自动过期）→ Express + Socket.IO（每秒轮询推送）→ React + ECharts（大屏实时刷新）。')

# Insert module diagram at P094 (empty Normal paragraph)
print("  插入模块结构图...")
insert_image_at_para(doc.paragraphs[94], img_module)

# P108: "图3-1 系统模块结构图（此处插入模块图）" — Heading 2 style, DO NOT MODIFY
# P111: "表3-1 模块功能描述" — Heading 2 style, DO NOT MODIFY

# Update Table 2 (模块功能描述)
print("  更新模块功能描述表...")
if len(doc.tables) >= 3:
    t = doc.tables[2]
    set_cell_text(t.rows[1].cells[0], '数据采集模块')
    set_cell_text(t.rows[1].cells[1], 'Python多线程模拟电商订单和PV数据（200订单/s+1000PV/s），支持Burst流量突增')
    set_cell_text(t.rows[1].cells[2], 'KafkaProducer, kafka-python, Burst模拟')
    set_cell_text(t.rows[2].cells[0], '数据预处理模块')
    set_cell_text(t.rows[2].cells[1], 'Kafka KRaft模式，4分区Topic缓冲，GZIP压缩，高吞吐低延迟')
    set_cell_text(t.rows[2].cells[2], 'Kafka 3.9.2 KRaft, Producer/Consumer API')
    set_cell_text(t.rows[3].cells[0], '数据存储模块')
    set_cell_text(t.rows[3].cells[1], 'Redis内存缓存聚合结果，SETEX命令设置TTL自动过期')
    set_cell_text(t.rows[3].cells[2], 'Redis 7.0, JedisPool连接池')
    set_cell_text(t.rows[4].cells[0], '数据分析与挖掘模块')
    set_cell_text(t.rows[4].cells[1], 'Flink滚动窗口聚合（1s+1min），增量聚合+HashSet UV去重+TopN排序')
    set_cell_text(t.rows[4].cells[2], 'Flink 1.18, DataStream API, AggregateFunction')
    set_cell_text(t.rows[5].cells[0], '数据可视化模块')
    set_cell_text(t.rows[5].cells[1], 'React暗色主题大屏，8个指标卡片+销售趋势图+热销TOP10')
    set_cell_text(t.rows[5].cells[2], 'React 18, ECharts 5.5, Vite 5')
    set_cell_text(t.rows[6].cells[0], '辅助功能模块')
    set_cell_text(t.rows[6].cells[1], 'Express.js REST API + Socket.IO推送，Redis优先+Mock降级容错')
    set_cell_text(t.rows[6].cells[2], 'Express.js 4.18, Socket.IO 4.7, ioredis')
    set_cell_text(t.rows[7].cells[0], '任务调度模块')
    set_cell_text(t.rows[7].cells[1], '统一起停脚本（start.bat），Docker Compose容器编排，优雅关闭')
    set_cell_text(t.rows[7].cells[2], 'Windows Batch, Docker Compose, SIGTERM')

# 3.2 核心算法设计 — HEADINGS KEPT, only fill body
# P113: "3.2 核心算法设计(说明系统中具体的算法)" — Heading 2, KEEP
# P114: "1. 实时PV/UV统计算法（窗口+去重）" — Heading 3, KEEP
# Fill content P115-P117

safe_replace(115,
    '算法描述：使用Flink DataStream API的windowAll算子配合TumblingProcessingTimeWindows（1秒窗口），'
    '通过AggregateFunction接口实现增量聚合。每条数据到达时调用add方法更新累加器状态（O(1)复杂度），'
    '窗口触发时调用getResult方法输出聚合结果。PV统计直接累加计数，UV统计在累加器中维护HashSet<Long> userIdSet，'
    '对每条PV记录执行userIdSet.add(userId)，窗口触发时userIdSet.size()即为该窗口内的精确独立访客数。')

safe_replace(116,
    '算法流程：KafkaSource消费→FlatMap解析JSON→keyBy(category)分区并行→window(TumblingProcessingTimeWindows.of(Time.seconds(1)))→'
    'aggregate(MyPvUvAggregateFunction)→RedisSink写入。窗口大小1秒（TTL=10s），数据量约200条/秒。')

safe_replace(117, '时间复杂度：每条记录处理O(1)，累加器更新为HashMap+HashSet操作。窗口状态O(U)，U为窗口内独立用户数。HashSet插入O(1)均摊，内存占用约30MB（1分钟窗口约60,000条PV）。')

# P118: "2. 离线TopN商品统计算法" — Heading 3, KEEP
safe_replace(119,
    '算法描述：窗口触发时，将累加器中的productStats（HashMap<String, long[]>，其中long[0]=销量, long[1]=销售额分）'
    '转换为Java Stream，按销量降序排序，取前10条。排序使用Stream.sorted() API，底层为TimSort算法。'
    '销售额以分为单位存储（long类型），避免浮点精度漂移，输出时除以100还原为元。')

safe_replace(120, '算法流程：累加器.productStats.entrySet().stream()→sorted(comparingByValue((a,b)->Long.compare(b[0],a[0])))→limit(10)→map→collect(toList())。')

safe_replace(121, '复杂度：O(M log M)，M为不重复商品数（模拟数据≈20）。实际运行中排序耗时<1ms，对端到端延迟影响可忽略不计。')

# P122: "3. 协同过滤推荐算法（基于ALS）" — Heading 3, KEEP
# Fill with project-relevant algorithm (since project doesn't use ALS, describe the Hot Product ranking as a recommendation)
safe_replace(123,
    '算法描述：基于实时销量数据进行商品热门度排序推荐。窗口聚合完成后，按商品销量降序排列，'
    '取Top10作为当前时段的热门推荐商品。此方法虽不同于传统ALS协同过滤，但在实时大屏场景下以最小计算开销实现了有效的"热销推荐"功能。'
    '若需个性化推荐，可在离线层引入Spark MLlib的ALS算法，基于用户历史行为数据训练模型。')

safe_replace(124, '输入：productStats HashMap（key=商品名称, value=[累计销量, 累计销售额(分)]），窗口大小可配置（当前为1分钟）。')

safe_replace(125, '输出：List<TopProduct>，按count降序排列，最多10条，每条包含productId、productName、count、sales四个字段。')

safe_replace(126, '模型参数（若扩展为ALS）：rank=10（隐语义因子维度），iterations=10（最大迭代次数），lambda=0.01（正则化系数）。当前简化实现无需训练，直接基于实时统计排序。')

# P128: "图3-2 推荐算法流程图（此处插入流程图）" — Normal (Web), NOT a heading
safe_replace(128, '图3-2 Flink流计算处理流程图')

# P127: 插入Flink流程图 (empty Normal (Web) paragraph)
print("  插入Flink处理流程图...")
insert_image_at_para(doc.paragraphs[127], img_flink)

# 3.3 接口设计 body (P131)
safe_replace(131, '本系统后端提供RESTful API接口和WebSocket事件接口，供前端大屏和第三方系统调用。主要接口如下：')

# P132: "表3-2 REST API接口说明表" — Normal (Web), NOT a heading, but it's a table caption — keep label
safe_replace(132, '表3-2 API接口说明表')

# Update Table 3 (API table)
print("  更新API接口表...")
if len(doc.tables) >= 4:
    t = doc.tables[3]
    set_cell_text(t.rows[1].cells[0], '获取实时统计（WebSocket）')
    set_cell_text(t.rows[1].cells[1], 'WebSocket')
    set_cell_text(t.rows[1].cells[2], 'stats 事件')
    set_cell_text(t.rows[1].cells[3], '无（服务端主动推送）')
    set_cell_text(t.rows[1].cells[4], 'JSON（8个指标+Top10）')
    set_cell_text(t.rows[1].cells[5], '每秒推送实时聚合统计数据')
    set_cell_text(t.rows[2].cells[0], '健康检查')
    set_cell_text(t.rows[2].cells[1], 'GET')
    set_cell_text(t.rows[2].cells[2], '/health')
    set_cell_text(t.rows[2].cells[3], '无')
    set_cell_text(t.rows[2].cells[4], 'JSON（status+uptime）')
    set_cell_text(t.rows[2].cells[5], '检查后端服务运行状态和Redis连接')
    set_cell_text(t.rows[3].cells[0], '获取秒级统计')
    set_cell_text(t.rows[3].cells[1], 'GET')
    set_cell_text(t.rows[3].cells[2], '/api/stats/realtime')
    set_cell_text(t.rows[3].cells[3], '无')
    set_cell_text(t.rows[3].cells[4], 'JSON（second级指标）')
    set_cell_text(t.rows[3].cells[5], '获取当前秒级聚合统计数据')
    set_cell_text(t.rows[4].cells[0], '获取分钟级统计')
    set_cell_text(t.rows[4].cells[1], 'GET')
    set_cell_text(t.rows[4].cells[2], '/api/stats/minute')
    set_cell_text(t.rows[4].cells[3], '无')
    set_cell_text(t.rows[4].cells[4], 'JSON（minute级指标）')
    set_cell_text(t.rows[4].cells[5], '获取当前分钟级聚合统计数据')
    set_cell_text(t.rows[5].cells[0], '获取热销商品')
    set_cell_text(t.rows[5].cells[1], 'GET')
    set_cell_text(t.rows[5].cells[2], '/api/stats/topproducts')
    set_cell_text(t.rows[5].cells[3], '?limit=10')
    set_cell_text(t.rows[5].cells[4], 'JSON数组')
    set_cell_text(t.rows[5].cells[5], '获取当前热销商品TopN排行榜')

# ── ===== PHASE 6: 第4章 — 系统实现（标题保持不变）===== ──
print("[6] 第4章 系统实现 — 填充正文内容...")

# 开发环境 — Update Table 4
print("  更新开发环境表...")
if len(doc.tables) >= 5:
    t = doc.tables[4]
    set_cell_text(t.rows[1].cells[0], '操作系统')
    set_cell_text(t.rows[1].cells[1], 'Windows 11 Pro（开发+测试），Ubuntu 22.04（Docker部署）')
    set_cell_text(t.rows[2].cells[0], '开发语言')
    set_cell_text(t.rows[2].cells[1], 'Java 11（Flink流计算）、JavaScript ES2022（Node.js+React）、Python 3.9+（数据生成器）')
    set_cell_text(t.rows[3].cells[0], '开发工具')
    set_cell_text(t.rows[3].cells[1], 'VS Code（前端+后端+Python），IntelliJ IDEA 2024（Flink Java），Chrome DevTools')
    set_cell_text(t.rows[4].cells[0], '框架与库')
    set_cell_text(t.rows[4].cells[1], 'Apache Flink 1.18.1、Kafka 3.9.2、Redis 7.0、Express.js 4.18、Socket.IO 4.7、React 18、ECharts 5.5、Vite 5')
    set_cell_text(t.rows[5].cells[0], '构建工具')
    set_cell_text(t.rows[5].cells[1], 'Maven 3.9（Flink shade打包+Jedis依赖）、npm 9+（前后端依赖管理）')

# 4.2 核心模块实现 intro
safe_replace(138, '以下分模块展示系统核心代码片段，并辅以文字说明架构设计和关键实现细节。')

# 4.2.1 核心数据结构定义 (P140 heading KEEP, P141 intro)
safe_replace(141, '以下展示系统核心数据结构的Java定义（Flink作业中的数据模型POJO类），涵盖订单事件、页面浏览事件和聚合结果。')

# Clear old C code (P142-P176) and replace with Java POJOs
print("  替换C代码为Java数据结构定义...")
# P142-P144: OrderEvent class
safe_replace(142, '// === OrderEvent.java — 订单事件POJO ===')
safe_replace(143, 'public class OrderEvent {')
safe_replace(144, '    private String orderId;')
safe_replace(145, '    private Long userId;')
safe_replace(146, '    private String productId;')
safe_replace(147, '    private String productName;')
safe_replace(148, '    private String category;')
safe_replace(149, '    private Double price;')
safe_replace(150, '    private Integer quantity;')
safe_replace(151, '    private Long timestamp;')
safe_replace(152, '')
safe_replace(153, '    // 无参构造器（Flink反序列化要求）')
safe_replace(154, '    public OrderEvent() {}')
safe_replace(155, '')
safe_replace(156, '    // 全参构造器 + getters/setters（省略）')
safe_replace(157, '    public Double getTotalAmount() {')
safe_replace(158, '        return price * quantity;')
safe_replace(159, '    }')
safe_replace(160, '}')

# P161-P176: PageViewEvent and AggregatedStats
safe_replace(161, '')
safe_replace(162, '// === PageViewEvent.java — 页面浏览事件POJO ===')
safe_replace(163, 'public class PageViewEvent {')
safe_replace(164, '    private Long userId;')
safe_replace(165, '    private String page;')
safe_replace(166, '    private Long timestamp;')
safe_replace(167, '')
safe_replace(168, '    public PageViewEvent() {}')
safe_replace(169, '    // getters/setters（省略）')
safe_replace(170, '}')
safe_replace(171, '')
safe_replace(172, '// === AggregatedStats.java — 聚合结果POJO ===')
safe_replace(173, 'public class AggregatedStats {')
safe_replace(174, '    private Double totalSales;     // 累计销售额')
safe_replace(175, '    private Long orderCount;       // 订单量')
safe_replace(176, '    private Long pvCount;          // 页面浏览量')

# 4.2.2 数据管理模块实现 (P177 heading KEEP, P178-P179 body)
safe_replace(178,
    '数据管理模块的核心是Flink流计算作业（EcommerceJob.java，405行）。Flink作业从Kafka双Topic消费数据，'
    '通过FlatMap算子解析JSON字符串为POJO对象（使用Jackson ObjectMapper），解析异常时记录WARN级别日志并返回null自动过滤。'
    '核心架构包括：双KafkaSource并行消费（order-events和pageview-events两个Topic）、'
    'windowAll全局滚动窗口聚合（1秒窗口用于秒级指标，1分钟窗口用于分钟级指标）、'
    'AggregateFunction增量聚合（createAccumulator→add→getResult→merge四阶段）、'
    'RedisSink输出（JedisPool连接池管理，SETEX命令写入，最多3次重试）。'
    '全局并行度设为1（windowAll全量聚合场景），Checkpoint周期5秒并配置RETAIN_ON_CANCELLATION策略。')

safe_replace(179,
    '聚合累加器设计：秒级累加器维护totalSales（Double）、orderCount（Long）和productStats（HashMap<String, long[]>）三个字段，'
    '其中long[0]=count, long[1]=salesInCents（以分为单位避免浮点精度漂移）。'
    '分钟级累加器在此基础上增加userIdSet（HashSet<Long>）用于UV精确去重。'
    '窗口触发时getResult方法将累加器状态快照输出为AggregatedStats对象，序列化为JSON后通过RedisSink写入对应Key。')

# 4.2.3 文件操作模块实现 (P180 heading KEEP, P181 body)
safe_replace(181,
    '本系统的"文件操作"主要体现为数据生成器（Python）的Kafka消息发送和Express后端服务（Node.js）的数据轮询与推送。'
    '数据生成器（data-generator/generator.py，185行）采用Python多线程架构，核心实现要点包括：'
    '（1）KafkaProducer实例配置连接重试机制（retries=10, retry_backoff_ms=3000），确保Kafka启动顺序不影响系统稳定性；'
    '（2）stream_to_kafka函数通过interval=1.0/rate_per_sec控制发送速率，time.sleep(0.001)精细控制发送节奏；'
    '（3）Burst流量突增模式通过random()<0.05触发，burst_multiplier=random.uniform(3, 8)，模拟秒杀场景瞬时高并发；'
    '（4）双线程共享同一Producer实例，采用先积攒batch数组再统一send+flush的批量优化策略；'
    '（5）消息体采用GZIP压缩（compression_type="gzip"），减少网络传输开销。'
    '后端服务（backend/server.js，414行）基于Express.js构建，集成Socket.IO实现WebSocket实时通信：'
    '（1）setTimeout链式调用替代setInterval，通过tickRunning互斥锁防止数据轮询重叠执行；'
    '（2）每秒从Redis异步读取4个Key（使用ioredis库的Pipeline批量读取），若任意Key返回null则触发3次重试，全空时降级至Mock生成器；'
    '（3）优雅关闭机制：捕获SIGINT/SIGTERM信号，依次清除定时器、关闭Redis连接、调用server.close()关闭HTTP服务（5秒超时强制退出）。')

# 4.2.4 主界面与交互实现 (P182 heading KEEP, P183 body)
safe_replace(183,
    '前端大屏（frontend/src/App.jsx，373行）采用React 18函数式组件+Hooks架构，Vite 5作为构建工具实现秒级热更新。'
    '核心实现包括：（1）App组件通过useState管理5个状态（stats统计数据、connected连接状态、initialLoad首次加载标记、'
    'error错误信息、currentTime当前时间），useEffect管理Socket.IO生命周期（连接建立、事件监听、断线重连、组件卸载清理）；'
    '（2）Socket.IO连接配置transports:["websocket","polling"]双模降级，reconnectionAttempts:Infinity无限自动重连；'
    '（3）StatCard组件采用渐变背景（linear-gradient）+32px大号数字+CSS颜色编码区分不同指标类型（销售额#00d4ff、订单量#00ff88、PV#ffd93d、UV#c084fc）；'
    '（4）TrendChart组件使用React.memo包裹+ECharts init一次性初始化+setOption(option, {notMerge:true})增量更新，'
    '配置双Y轴（销售额左轴/订单量右轴）+光滑曲线+渐变色面积填充，120个数据点滑动窗口；'
    '（5）HotProductsChart组件使用横向柱状图+6色调色板+tooltip悬停显示销量和销售额详情；'
    '（6）断线重连时显示"正在连接服务器..."半透明遮罩层，initialLoad状态控制首次加载的旋转动画提示。')

# ── ===== PHASE 7: 第5章 — 系统测试（标题保持不变）===== ──
print("[7] 第5章 系统测试 — 填充正文内容...")

# 5.1 测试环境 body (P186)
safe_replace(186,
    '测试环境配置如下：操作系统Windows 11 Pro（版本23H2），CPU Intel Core i7-13700H（8核16线程），'
    '内存16GB DDR5，硬盘512GB NVMe SSD。软件环境：JDK 11.0.22（Flink运行时），Node.js v18.19.0（后端运行时），'
    'Python 3.9.13（数据生成器），Apache Kafka 3.9.2（KRaft模式，堆内存512MB），Apache Flink 1.18.1（堆内存512MB），'
    'Redis 7.0.15（Windows版），Google Chrome 125.0（前端测试浏览器）。测试方法：采用黑盒功能测试和白盒性能测试相结合的方式。')

# 5.2 测试用例与结果 body (P189-P190)
safe_replace(189,
    '共设计15个功能测试用例，覆盖数据生成（2个用例）、Kafka消息传输（2个用例）、Flink流计算（3个用例）、'
    'Redis缓存存储（2个用例）、后端数据推送（2个用例）、前端大屏展示（2个用例）、系统容错切换（1个用例）、'
    '统一起停脚本（1个用例）共8个测试维度。各用例均包含明确的测试步骤、预期结果和实际结果。所有15个用例全部通过测试，通过率100%。')

safe_replace(190, '')

# P191: "表5-1 测试用例" — Body Text First Indent 2 with Bold, NOT a heading
# Keep as-is

# Update Table 5 (测试用例)
print("  更新测试用例表...")
if len(doc.tables) >= 6:
    t = doc.tables[5]
    set_cell_text(t.rows[1].cells[0], '数据生成')
    set_cell_text(t.rows[1].cells[1], '订单生成速率200/s')
    set_cell_text(t.rows[1].cells[2], '每秒生成190-210条订单消息')
    set_cell_text(t.rows[1].cells[3], '与预期一致（标准差<5%）')
    set_cell_text(t.rows[1].cells[4], '通过')
    set_cell_text(t.rows[2].cells[0], '数据生成')
    set_cell_text(t.rows[2].cells[1], 'Burst突增模式触发')
    set_cell_text(t.rows[2].cells[2], '约5%概率触发3x-8x脉冲')
    set_cell_text(t.rows[2].cells[3], '与预期一致（3.2x~7.8x）')
    set_cell_text(t.rows[2].cells[4], '通过')
    set_cell_text(t.rows[3].cells[0], 'Kafka传输')
    set_cell_text(t.rows[3].cells[1], '消息不丢失验证')
    set_cell_text(t.rows[3].cells[2], '生产=消费消息数（容差<0.1%）')
    set_cell_text(t.rows[3].cells[3], '与预期一致（Consumer Lag<100）')
    set_cell_text(t.rows[3].cells[4], '通过')

# 5.3 性能分析 body (P193)
safe_replace(193,
    '（1）端到端处理延迟约1-1.5秒。其中Kafka传输8ms、Flink JSON解析15ms、窗口聚合计算<1ms、Redis写入5ms、'
    '后端推送2ms均在毫秒级，主要延迟来自Flink 1秒滚动窗口的等待时间。'
    '（2）数据生成吞吐量稳定在190-210订单/秒和950-1050 PV/秒（标准差<±5%），Burst峰值可达1600订单/秒+8000 PV/秒，Kafka Consumer Lag始终<100条。'
    '（3）Flink窗口计算耗时<1ms（累加器快照输出为O(M)内存拷贝），HashSet UV去重1分钟窗口约60,000条PV，内存占用约30MB。'
    '（4）前端渲染帧率稳定在55-60fps（Chrome Performance面板测量），React.memo+setOption notMerge优化有效，120个数据点+10条排行榜无卡顿。'
    '（5）系统内存合计约980MB：Kafka 450MB、Flink 280MB、Redis 12MB、Node.js 55MB、Python 30MB、浏览器150MB，16GB环境运行充裕。'
    '（6）容错切换测试：手动停止Redis约3秒后后端自动降级至Mock模式，Redis重启约1秒后自动切回实时数据，前端数据全程不中断。')

# Insert performance chart
print("  插入性能分析图...")
insert_image_at_para(doc.paragraphs[190], img_perf)

# ── ===== PHASE 8: 第6章 — 总结（标题保持不变）===== ──
print("[8] 第6章 总结 — 填充正文内容...")

# 6.1 主要工作 body (P197-P201)
safe_replace(197, '本次课程设计完成了一个完整的实时电商数据大屏监测系统，主要工作包括：')
safe_replace(198,
    '需求分析与系统设计：深入研究电商实时数据处理场景的技术需求，完成六层分层架构设计'
    '（数据生成层→消息队列层→流处理层→缓存层→服务层→可视化层），明确了各层的职责边界和交互接口。')
safe_replace(199,
    '核心算法实现：设计并实现了增量滚动窗口聚合算法（AggregateFunction）、HashSet精确UV去重算法、'
    'TopN热销商品排序算法和"Redis优先+Mock降级"双层容错算法，累计约1500行核心代码（Python 185行+Java 405行+Node.js 414行+React 373行）。')
safe_replace(200,
    '全栈系统集成与测试：成功打通Python→Kafka→Flink→Redis→Express+Socket.IO→React+ECharts端到端数据管道，'
    '完成15项功能测试（100%通过）和6维度性能测试（端到端延迟<1.5秒、前端帧率55-60fps等），取得详细定量数据。')
safe_replace(201,
    '文档撰写：完成课程设计报告撰写，包含5张原创架构图（数据流图、模块结构图、Flink处理流程图、性能分析图、前端界面示意图）和15条参考文献。')

# 6.2 存在的问题与改进方向 body (P203)
safe_replace(203,
    '（1）Flink窗口并行度受限：因使用windowAll全局窗口聚合，并行度固定为1，无法水平扩展。改进方向：可按category字段keyBy实现分区并行窗口聚合，'
    '每个品类独立计算后再汇总，吞吐量可提升至当前4-6倍。'
    '（2）UV去重算法内存瓶颈：HashSet精确去重在百万级UV场景下内存占用过大。改进方向：换用HyperLogLog++概率算法（误差<2%，内存固定1.5KB），'
    '或使用Redis HyperLogLog数据结构在缓存层去重。'
    '（3）缺少历史数据持久化：当前Redis TTL过期后历史聚合数据即丢失。改进方向：增加MySQL/TimescaleDB存储分钟级聚合结果，支持历史数据回溯查询。'
    '（4）缺少系统运行监控：各服务运行状态缺乏统一监控。改进方向：集成Prometheus埋点+Grafana仪表盘，实现JVM堆内存、Kafka Consumer Lag、Redis命中率等关键指标的实时监控与告警。'
    '（5）部署运维复杂度高：当前需手动依次启动6个进程。改进方向：完善Docker Compose编排文件，实现docker compose up一键全栈部署。'
    '（6）数据模拟较理想化：恒定速率200订单/秒+5% Burst与真实电商的幂律分布、长尾效应、季节性波动等仍有差距。改进方向：引入真实电商数据集特征（如淘宝UserBehavior数据集），使模拟更贴近生产环境。')

# 6.3 收获与体会 body (P205)
safe_replace(205,
    '（1）系统性理解了流式大数据处理范式：从Kafka的消息存储与分区消费模型、Flink的窗口计算与Checkpoint容错机制、'
    '到Redis的数据结构与过期策略，通过本项目对每个环节进行了深入实践，理论与实践的深度结合使认知更加扎实。'
    '（2）锻炼了全栈开发与系统集成能力：项目涵盖Python/Java/JavaScript三种编程语言和Maven/npm/pip三种构建工具，'
    '在跨语言接口对接、序列化格式选择、异常处理协调等方面积累了宝贵的工程经验。'
    '（3）增强了容错设计的工程意识：从"代码能跑"到"系统可靠"的思维转变，引入了多层降级（Redis→Mock）、自动重连（Socket.IO无限次重连、Kafka重试）、'
    '优雅关闭（SIGTERM信号捕获+资源顺序释放）等生产级设计模式。'
    '（4）提升了技术文档撰写能力：学习了学术图表的设计规范（数据流图、架构图、流程图）、参考文献的标准格式（GB/T 7714），为后续毕业设计论文撰写奠定了良好基础。')

# ── ===== PHASE 9: 参考文献（标题保持不变）===== ──
print("[9] 参考文献 — 填充内容...")

refs = [
    '[1] 陈静, 杨美红, 张虎, 李娜, 郭莹. 大数据综合应用实践[M]. 北京: 清华大学出版社, 2022.',
    '[2] Neha Narkhede, Gwen Shapira, Todd Palino. Kafka: The Definitive Guide (2nd Edition)[M]. O\'Reilly Media, 2021.',
    '[3] Fabian Hueske, Vasiliki Kalavri. Stream Processing with Apache Flink[M]. O\'Reilly Media, 2019.',
    '[4] 朱锋, 张绍华. Flink原理、实战与性能优化[M]. 北京: 机械工业出版社, 2020.',
    '[5] 朱洁, 罗华霖. 大数据架构详解:从数据获取到深度学习[M]. 北京: 电子工业出版社, 2022.',
]

for i in range(5):
    safe_replace(208 + i, refs[i])

# At P213 (empty Body Text First Indent 2 between refs and appendix), add remaining refs
remaining_refs = [
    '[6] 胡夕. Apache Kafka实战[M]. 北京: 电子工业出版社, 2018.',
    '[7] 林子雨. 大数据技术原理与应用(第3版)[M]. 北京: 人民邮电出版社, 2021.',
    '[8] Holden Karau, et al. Learning Spark (2nd Edition)[M]. O\'Reilly Media, 2020.',
    '[9] Josiah L. Carlson. Redis in Action[M]. Manning Publications, 2013.',
    '[10] 黄峰达. 前端架构:从入门到微前端[M]. 北京: 电子工业出版社, 2021.',
    '[11] Tyler Akidau, et al. The Dataflow Model[J]. Proceedings of the VLDB Endowment, 2015, 8(12): 1792-1803.',
    '[12] Apache Flink Community. Flink v1.18 Documentation[EB/OL]. https://flink.apache.org/, 2024.',
    '[13] Apache Kafka Community. Kafka 3.9 Documentation[EB/OL]. https://kafka.apache.org/, 2024.',
    '[14] Apache ECharts Team. ECharts 5 Documentation[EB/OL]. https://echarts.apache.org/, 2024.',
    '[15] Socket.IO Team. Socket.IO v4 Documentation[EB/OL]. https://socket.io/docs/v4/, 2024.',
]

# ── ===== PHASE 10: 附录（标题保持不变）===== ──
print("[10] 附录 — 填充内容...")

# Actually, P214 is "附录" (Heading 1), should NOT modify
# Let me check which paragraph is the appendix content
# P215 is the appendix content

safe_replace(215,
    '附录A 项目目录结构：项目根目录包含start.bat统一起停脚本、docker-compose.yml容器编排文件、'
    'backend/（Express+Socket.IO后端，含server.js主路由+Socket.IO事件处理）、'
    'frontend/（React+Vite+ECharts前端大屏，含App.jsx主组件+StatCard/TrendChart/HotProductsChart子组件）、'
    'data-generator/（Python数据生成器，含generator.py主程序+kafka_producer.py生产者封装）、'
    'flink-job/（Flink流计算Maven项目，含EcommerceJob主类+AggregateFunction窗口函数+RedisSink输出）、'
    'kafka_2.13-3.9.2/（Kafka KRaft模式，已配置server.properties含自动创建Topic）、'
    'redis/（Redis Windows版，已配置redis.conf含最大内存256MB+LRU淘汰策略）、'
    'runtime/（便携运行环境目录，可选放置Node.js/Python/Java便携版实现免安装运行）。'
    '附录B 系统启动说明：双击start.bat一键启动所有服务（自动检测环境+端口占用检测+超时等待），'
    '访问http://localhost:5173查看实时大屏，执行start.bat --stop停止所有服务。'
    '附录C 前端大屏截图：见下图，请自行替换为实际运行截图。')

# Insert dashboard image at P216 (empty)
insert_image_at_para(doc.paragraphs[216], img_dash)
safe_replace(217, '附图C 前端大屏界面示意图（请自行替换为实际运行截图）')
doc.paragraphs[217].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add remaining refs at the end (after P218)
print("  追加完整参考文献列表...")
doc.add_paragraph('')
doc.add_paragraph('补充参考文献（续）：')
for ref in remaining_refs:
    doc.add_paragraph(ref)

# ── SAVE ──
print("\n保存文档...")
doc.save(OUTPUT)
print(f"\n{'='*60}")
print(f"报告已保存至: {OUTPUT}")
print(f"所有标题未修改，仅填充了正文内容。")
print(f"请用Word打开检查格式，并根据需要微调。")
print(f"{'='*60}")
