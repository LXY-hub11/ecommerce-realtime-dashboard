# -*- coding: utf-8 -*-
"""Verify the generated report — check pages 1-2 intact, content filled, formatting preserved"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

OUTPUT = r"D:\项目\1\电商大屏监测系统\实时电商数据大屏监测系统_课程设计报告.docx"
doc = Document(OUTPUT)

print("=" * 80)
print("VERIFICATION: Generated Report Analysis")
print("=" * 80)

# Check pages 1-2 (P000-P019) are unchanged
print("\n--- Pages 1-2 (P000-P019): Should be UNCHANGED ---")
for i in range(20):
    para = doc.paragraphs[i]
    text = para.text[:80] if para.text else '(empty)'
    style = para.style.name if para.style else 'None'
    print(f"  P{i:03d} [{style}] {text}")

# Check headings
print("\n--- Headings: Should be PRESERVED ---")
heading_count = 0
for i, para in enumerate(doc.paragraphs):
    style_name = para.style.name if para.style else ''
    if 'Heading' in style_name:
        heading_count += 1
        print(f"  P{i:03d} [{style_name}] {para.text[:80]}")
print(f"  Total headings: {heading_count}")

# Check key content sections
print("\n--- Key Content Verification ---")
checks = {
    21: '摘要内容',
    24: 'Heading: 1 系统概述',
    55: 'Heading: 2 系统分析',
    91: 'Heading: 3 系统设计',
    134: 'Heading: 4 系统实现',
    184: 'Heading: 5 系统测试',
    195: 'Heading: 6 总结',
    207: 'Heading: 参考文献',
    214: 'Heading: 附录',
}
for idx, desc in checks.items():
    para = doc.paragraphs[idx]
    text = para.text[:100] if para.text else '(empty)'
    print(f"  P{idx:03d} ({desc}): {text}")

# Check images
print("\n--- Image Check ---")
img_count = 0
for para in doc.paragraphs:
    for run in para.runs:
        if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
            img_count += 1
        # Also check inline shapes
        drawings = run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
        img_count += len(drawings)
print(f"  Images found: {img_count}")

# Check tables
print("\n--- Tables Check ---")
for t_idx, table in enumerate(doc.tables):
    print(f"  Table {t_idx}: {len(table.rows)} rows × {len(table.columns)} cols")
    print(f"    Header: {[cell.text[:40] for cell in table.rows[0].cells]}")

# Check total paragraphs
print(f"\n--- Summary ---")
print(f"  Total paragraphs: {len(doc.paragraphs)}")
print(f"  Total tables: {len(doc.tables)}")

print("\nVerification complete!")
