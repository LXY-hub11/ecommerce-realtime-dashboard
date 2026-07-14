# -*- coding: utf-8 -*-
"""Analyze template structure — paragraphs, headings, tables, styles"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

TEMPLATE = r"D:\项目\1\电商大屏监测系统\课程设计报告模板-大数据综合实训-0.docx"
doc = Document(TEMPLATE)

print("=" * 80)
print("PARAGRAPH ANALYSIS (index, style, first 100 chars of text)")
print("=" * 80)

for i, para in enumerate(doc.paragraphs):
    style_name = para.style.name if para.style else 'None'
    text_preview = para.text[:100] if para.text else '(empty)'
    # Check for headings
    is_h = 'Heading' in style_name or 'heading' in style_name or '标题' in style_name
    marker = ' ★HEADING★' if is_h else ''
    # Check alignment
    align = str(para.alignment) if para.alignment else 'None'
    # Check font info from first run
    font_info = ''
    if para.runs:
        run = para.runs[0]
        font_info = f'font={run.font.name}, size={run.font.size}, bold={run.font.bold}'

    print(f'P{i:03d} [{style_name}]{marker} align={align} | {font_info}')
    print(f'      → {text_preview}')

print("\n" + "=" * 80)
print("TABLE ANALYSIS")
print("=" * 80)

for t_idx, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns)
    print(f'\nTable {t_idx}: {rows} rows × {cols} columns')
    for r_idx, row in enumerate(table.rows):
        cells_text = []
        for c_idx, cell in enumerate(row.cells):
            cell_text = cell.text[:60] if cell.text else '(empty)'
            cells_text.append(f'[{c_idx}] {cell_text}')
        print(f'  Row {r_idx}: {" | ".join(cells_text)}')

print("\n" + "=" * 80)
print("SECTION / PAGE BREAKS")
print("=" * 80)
for i, para in enumerate(doc.paragraphs):
    for run in para.runs:
        if run._element.xml.count('w:br w:type="page"') > 0 or 'pageBreak' in str(run._element.xml):
            print(f'P{i:03d}: PAGE BREAK found')

print("\nDone.")
