# -*- coding: utf-8 -*-
"""Detailed format comparison between template and generated report"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu

TEMPLATE = r"D:\项目\1\电商大屏监测系统\课程设计报告模板-大数据综合实训-0.docx"
OUTPUT = r"D:\项目\1\电商大屏监测系统\实时电商数据大屏监测系统_课程设计报告.docx"

doc_t = Document(TEMPLATE)
doc_o = Document(OUTPUT)

def para_info(para):
    """Extract detailed formatting info from a paragraph"""
    style = para.style.name if para.style else 'None'
    pf = para.paragraph_format
    align = str(para.alignment) if para.alignment else 'None'

    # Indentation
    left_indent = pf.left_indent
    first_line_indent = pf.first_line_indent
    right_indent = pf.right_indent

    # Spacing
    space_before = pf.space_before
    space_after = pf.space_after
    line_spacing = pf.line_spacing

    # Run info (first run)
    r_font = None; r_size = None; r_bold = None; r_italic = None
    if para.runs:
        r = para.runs[0]
        r_font = r.font.name
        r_size = r.font.size
        r_bold = r.font.bold
        r_italic = r.font.italic

    return {
        'style': style,
        'align': align,
        'left_indent': left_indent,
        'first_line_indent': first_line_indent,
        'right_indent': right_indent,
        'space_before': space_before,
        'space_after': space_after,
        'line_spacing': line_spacing,
        'font': r_font,
        'size': r_size,
        'bold': r_bold,
        'italic': r_italic,
    }

def fmt_emu(val):
    """Format EMU value to readable string"""
    if val is None:
        return 'None'
    emu = int(val)
    pt = emu / 12700
    cm = emu / 360000
    return f'{emu} EMU ({pt:.1f}pt, {cm:.2f}cm)'

def fmt_line_spacing(val):
    if val is None:
        return 'None'
    if isinstance(val, float):
        return f'{val:.1f}x'
    return f'{int(val)} EMU ({int(val)/12700:.1f}pt)'

print("=" * 120)
print("DETAILED FORMAT COMPARISON: Template vs Generated Report (P020 onwards)")
print("=" * 120)
print(f"{'Idx':<5} {'Style':<28} {'Align':<12} {'Font':<20} {'Size':<22} {'Bold':<6} {'1stLineIndent':<22} {'LeftIndent':<20}")
print("-" * 120)

format_issues = []

for i in range(20, max(len(doc_t.paragraphs), len(doc_o.paragraphs))):
    if i >= len(doc_t.paragraphs):
        print(f"P{i:03d} [EXTRA in output only]")
        continue
    if i >= len(doc_o.paragraphs):
        print(f"P{i:03d} [MISSING in output]")
        continue

    pt = doc_t.paragraphs[i]
    po = doc_o.paragraphs[i]
    it = para_info(pt)
    io = para_info(po)

    # Check for differences
    issues = []

    # Style must match
    if it['style'] != io['style']:
        issues.append(f"STYLE: {it['style']} → {io['style']}")

    # Font
    if it['font'] != io['font']:
        issues.append(f"FONT: {it['font']} → {io['font']}")

    # Size
    if it['size'] != io['size']:
        issues.append(f"SIZE: {fmt_emu(it['size'])} → {fmt_emu(io['size'])}")

    # Bold
    if it['bold'] != io['bold']:
        issues.append(f"BOLD: {it['bold']} → {io['bold']}")

    # First line indent
    if it['first_line_indent'] != io['first_line_indent']:
        issues.append(f"1stINDENT: {fmt_emu(it['first_line_indent'])} → {fmt_emu(io['first_line_indent'])}")

    # Left indent
    if it['left_indent'] != io['left_indent']:
        issues.append(f"LEFT: {fmt_emu(it['left_indent'])} → {fmt_emu(io['left_indent'])}")

    # Alignment
    if it['align'] != io['align']:
        issues.append(f"ALIGN: {it['align']} → {io['align']}")

    # Line spacing
    if str(it['line_spacing']) != str(io['line_spacing']):
        issues.append(f"LINESPACE: {it['line_spacing']} → {io['line_spacing']}")

    # Space before
    if str(it['space_before']) != str(io['space_before']):
        issues.append(f"SPACE_BEFORE: {it['space_before']} → {io['space_before']}")

    # Space after
    if str(it['space_after']) != str(io['space_after']):
        issues.append(f"SPACE_AFTER: {it['space_after']} → {io['space_after']}")

    if issues:
        format_issues.append((i, issues))
        status = "⚠ MISMATCH"
    else:
        status = "✓"

    # Print summary line
    text_preview = po.text[:40] if po.text else '(empty)'
    print(f"P{i:03d} {status:<2} [{it['style'][:26]:<26}] align={str(it['align'])[-8:]:<8} "
          f"font={str(it['font'])[:18]:<18} size={fmt_emu(it['size']):<20} "
          f"bold={str(it['bold']):<5} 1st={fmt_emu(it['first_line_indent']):<20} "
          f"left={fmt_emu(it['left_indent']):<18} →\"{text_preview}\"")

# Summary of issues
print("\n" + "=" * 120)
print(f"FORMAT ISSUES FOUND: {len(format_issues)}")
print("=" * 120)
for idx, issues in format_issues:
    print(f"\nP{idx:03d}:")
    pt = doc_t.paragraphs[idx]
    po = doc_o.paragraphs[idx]
    print(f"  Template text: {pt.text[:80]}...")
    print(f"  Output text:   {po.text[:80]}...")
    for iss in issues:
        print(f"  → {iss}")

# Also check table formatting
print("\n" + "=" * 120)
print("TABLE FORMAT CHECK")
print("=" * 120)
for t_idx in range(min(len(doc_t.tables), len(doc_o.tables))):
    tt = doc_t.tables[t_idx]
    to = doc_o.tables[t_idx]
    print(f"\nTable {t_idx}: {len(tt.rows)} rows × {len(tt.columns)} cols")
    for r_idx in range(min(len(tt.rows), len(to.rows))):
        for c_idx in range(min(len(tt.rows[r_idx].cells), len(to.rows[r_idx].cells))):
            ct = tt.rows[r_idx].cells[c_idx]
            co = to.rows[r_idx].cells[c_idx]
            if ct.paragraphs and co.paragraphs:
                pt = ct.paragraphs[0]
                po_para = co.paragraphs[0]
                it = para_info(pt)
                io = para_info(po_para)
                if it['style'] != io['style'] or it['font'] != io['font'] or it['size'] != io['size']:
                    print(f"  [{r_idx},{c_idx}] STYLE: {it['style']}→{io['style']} | "
                          f"FONT: {it['font']}→{io['font']} | "
                          f"SIZE: {fmt_emu(it['size'])}→{fmt_emu(io['size'])} | "
                          f"BOLD: {it['bold']}→{io['bold']}")

print("\nDone.")
